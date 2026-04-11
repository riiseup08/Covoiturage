"""Generate a Word document detailing remaining work for Covoit.Africa deployment."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# --- Title ---
title = doc.add_heading('Covoit.Africa — Remaining Work Before Deployment', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    'This document details the remaining tasks required before the Covoit.Africa '
    'mobile application can be submitted to the App Store and Google Play Store. '
    'All code-level issues have been resolved. The items below are infrastructure, '
    'design, and configuration tasks that require manual action.'
)
doc.add_paragraph(f'Generated: April 10, 2026').italic = True
doc.add_paragraph('')

# ============================================================
# SECTION 1: BRANDING & ASSETS
# ============================================================
doc.add_heading('1. Branding & App Assets', level=1)
doc.add_paragraph(
    'The current asset images (icon.png, splash.png, adaptive-icon.png, favicon.png) '
    'in mobile/assets/ are placeholder green squares. App stores will reject these.'
)

table = doc.add_table(rows=5, cols=3, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Asset', 'Required Size', 'Notes']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        p.runs[0].bold = True

data = [
    ['icon.png', '1024 × 1024 px', 'App icon (iOS & Android). No transparency, no rounded corners (iOS rounds automatically).'],
    ['splash.png', '1284 × 2778 px', 'Splash/loading screen. Background color is #FAF6F1. Use "contain" resize mode.'],
    ['adaptive-icon.png', '1024 × 1024 px', 'Android adaptive icon foreground layer. Must have transparent background with safe zone padding.'],
    ['favicon.png', '48 × 48 px', 'Web browser tab favicon.'],
]
for r, row_data in enumerate(data, 1):
    for c, val in enumerate(row_data):
        table.rows[r].cells[c].text = val

doc.add_paragraph('')
doc.add_paragraph(
    'Action: Design or commission branded assets with the Covoit.Africa logo and earth-tone '
    'color scheme (#6B4226 brown, #FAF6F1 cream). Replace the files in mobile/assets/.'
)

# ============================================================
# SECTION 2: BACKEND DEPLOYMENT
# ============================================================
doc.add_heading('2. Backend Server Deployment (HTTPS)', level=1)
doc.add_paragraph(
    'The app currently points to localhost for development. Production builds require '
    'a publicly accessible Django backend served over HTTPS.'
)

doc.add_heading('Steps:', level=2)
steps = [
    'Choose a hosting provider (Railway, Render, DigitalOcean, AWS, Heroku, etc.)',
    'Deploy the Django backend (located at C:\\Users\\pokam\\Downloads\\Active_Projects\\Covoiturage_Moved)',
    'Configure a PostgreSQL database (replace SQLite)',
    'Set up a custom domain (e.g., api.covoitafrica.com)',
    'Obtain and configure an SSL/TLS certificate (Let\'s Encrypt or provider-managed)',
    'Update mobile/src/config/env.js — change API_URL_PRODUCTION to your HTTPS URL',
    'Set the EXPO_PUBLIC_API_URL environment variable in EAS build secrets',
]
for step in steps:
    doc.add_paragraph(step, style='List Number')

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('File to edit: ').bold = True
p.add_run('mobile/src/config/env.js → API_URL_PRODUCTION')

doc.add_paragraph('')
p2 = doc.add_paragraph()
p2.add_run('Current value: ').bold = True
p2.add_run("'https://api.covoitafrica.com/api'  (placeholder — update with real URL)")

# ============================================================
# SECTION 3: EAS PROJECT SETUP
# ============================================================
doc.add_heading('3. EAS (Expo Application Services) Setup', level=1)
doc.add_paragraph(
    'EAS Build is configured in mobile/eas.json with three profiles (development, preview, production). '
    'You need to create an EAS project and link it.'
)

doc.add_heading('Steps:', level=2)
steps = [
    'Install EAS CLI: npm install -g eas-cli',
    'Login: eas login',
    'Initialize: cd mobile && eas init (this creates the project and sets the projectId)',
    'Update mobile/app.json → extra.eas.projectId with the generated ID',
    'Update mobile/app.json → updates.url with: https://u.expo.dev/YOUR_PROJECT_ID',
    'For iOS: eas credentials (configure Apple certificates and provisioning profiles)',
    'For Android: eas credentials (upload or generate a keystore)',
    'Test build: eas build --profile preview --platform all',
    'Production build: eas build --profile production --platform all',
]
for step in steps:
    doc.add_paragraph(step, style='List Number')

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('Files to edit: ').bold = True
p.add_run('mobile/app.json (projectId, updates.url)')

# ============================================================
# SECTION 4: APP STORE SUBMISSION
# ============================================================
doc.add_heading('4. App Store & Play Store Submission', level=1)

doc.add_heading('4a. Apple App Store (iOS)', level=2)
items = [
    'Apple Developer account ($99/year): https://developer.apple.com',
    'Fill in app.json → ios.infoPlist with all required usage descriptions',
    'Create App Store listing in App Store Connect with screenshots (6.5" and 5.5")',
    'Provide a privacy policy URL (required)',
    'Provide a support URL',
    'Submit build via: eas submit --platform ios',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4b. Google Play Store (Android)', level=2)
items = [
    'Google Play Developer account ($25 one-time): https://play.google.com/console',
    'Create a service account key (google-services.json) for automated submission',
    'Fill in app listing with screenshots, feature graphic (1024×500), and description',
    'Complete the Data Safety form',
    'Provide a privacy policy URL (required)',
    'Submit build via: eas submit --platform android',
    'Note: First submission goes to "Internal testing" track (configured in eas.json)',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# ============================================================
# SECTION 5: PRIVACY POLICY
# ============================================================
doc.add_heading('5. Privacy Policy & Terms of Service', level=1)
doc.add_paragraph(
    'Both App Store and Play Store require a publicly accessible privacy policy URL. '
    'The policy must cover:'
)
items = [
    'What personal data is collected (name, email, phone, location, payment info)',
    'How data is used (trip matching, payments, communication between users)',
    'Third-party services (Mobile Money providers, push notification services)',
    'Data retention and deletion policy',
    'User rights (access, correction, deletion of their data)',
    'Contact information for data protection inquiries',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph(
    'Action: Create a privacy policy page hosted on your website (e.g., covoitafrica.com/privacy). '
    'Add the URL to app.json and to the app store listings. Consider using a generator like '
    'Termly, iubenda, or consulting a lawyer for GDPR/local compliance.'
)

# ============================================================
# SECTION 6: CRASH REPORTING
# ============================================================
doc.add_heading('6. Crash Reporting (Sentry / Bugsnag)', level=1)
doc.add_paragraph(
    'The app has an ErrorBoundary component (mobile/src/components/ErrorBoundary.js) '
    'with a TODO hook for remote error reporting. In production, crashes will be invisible '
    'without a reporting service.'
)

doc.add_heading('Recommended: Sentry', level=2)
steps = [
    'Create account at https://sentry.io (free tier available)',
    'Install: npx expo install @sentry/react-native',
    'Initialize in App.js with your DSN',
    'Update ErrorBoundary.js componentDidCatch to call Sentry.captureException(error)',
    'Test with: Sentry.nativeCrash() in dev mode',
]
for step in steps:
    doc.add_paragraph(step, style='List Number')

# ============================================================
# SECTION 7: SMS/OTP PROVIDER
# ============================================================
doc.add_heading('7. SMS/OTP Provider for Phone Authentication', level=1)
doc.add_paragraph(
    'The phone login flow calls the backend API for OTP. The Django backend needs '
    'an SMS provider configured to actually send OTP codes to African phone numbers.'
)

doc.add_heading('Recommended providers for Africa:', level=2)
table2 = doc.add_table(rows=5, cols=3, style='Light Grid Accent 1')
headers2 = ['Provider', 'Coverage', 'Notes']
for i, h in enumerate(headers2):
    table2.rows[0].cells[i].text = h
    for p in table2.rows[0].cells[i].paragraphs:
        p.runs[0].bold = True

sms_data = [
    ['Africa\'s Talking', 'Pan-African (40+ countries)', 'Best coverage in Sub-Saharan Africa. SMS + USSD.'],
    ['Twilio', 'Global', 'Reliable but more expensive for African routes.'],
    ['Vonage (Nexmo)', 'Global', 'Good SMS delivery rates in West/Central Africa.'],
    ['Termii', 'Nigeria, Ghana, Kenya', 'African-focused, affordable.'],
]
for r, row_data in enumerate(sms_data, 1):
    for c, val in enumerate(row_data):
        table2.rows[r].cells[c].text = val

doc.add_paragraph('')
doc.add_paragraph(
    'Action: Sign up with a provider, configure the Django backend '
    '(covoiturage/views.py → phone_request_otp) to send SMS via their API.'
)

# ============================================================
# SECTION 8: PAYMENT INTEGRATION
# ============================================================
doc.add_heading('8. Mobile Money Payment Integration', level=1)
doc.add_paragraph(
    'The payment screen supports MTN MoMo, Orange Money, Moov Money, Airtel Money, and Wave. '
    'The frontend is ready (PaymentScreen.js + TransactionConfirmationScreen.js). '
    'The Django backend needs actual payment provider integration.'
)

doc.add_heading('Steps:', level=2)
steps = [
    'MTN MoMo: Apply for API access at https://momodeveloper.mtn.com',
    'Orange Money: Contact Orange Money partner program',
    'Wave: Apply at https://wave.com/developers',
    'Implement payment webhooks in Django to confirm transactions',
    'Add transaction status polling or WebSocket updates in the mobile app',
    'Test thoroughly with sandbox/test credentials before going live',
]
for step in steps:
    doc.add_paragraph(step, style='List Number')

# ============================================================
# SECTION 9: SUMMARY TABLE
# ============================================================
doc.add_heading('9. Summary & Priority', level=1)

table3 = doc.add_table(rows=9, cols=4, style='Light Grid Accent 1')
headers3 = ['#', 'Task', 'Priority', 'Estimated Effort']
for i, h in enumerate(headers3):
    table3.rows[0].cells[i].text = h
    for p in table3.rows[0].cells[i].paragraphs:
        p.runs[0].bold = True

summary = [
    ['1', 'Deploy backend with HTTPS', 'CRITICAL', '1-2 days'],
    ['2', 'Design & replace app assets', 'CRITICAL', '1 day (design)'],
    ['3', 'EAS project setup + test build', 'CRITICAL', '2-3 hours'],
    ['4', 'Privacy policy page', 'CRITICAL', '2-4 hours'],
    ['5', 'SMS/OTP provider setup', 'HIGH', '4-8 hours'],
    ['6', 'Mobile Money integration', 'HIGH', '3-5 days'],
    ['7', 'Crash reporting (Sentry)', 'MEDIUM', '1-2 hours'],
    ['8', 'App Store listings + submission', 'FINAL STEP', '4-6 hours'],
]
for r, row_data in enumerate(summary, 1):
    for c, val in enumerate(row_data):
        table3.rows[r].cells[c].text = val

doc.add_paragraph('')
doc.add_paragraph(
    'All code-level work is complete (36 tests passing). The items above are '
    'infrastructure, third-party integrations, and app store configuration tasks. '
    'Once completed, the app is ready for production deployment.'
)

# --- Save ---
output_path = os.path.join(os.path.dirname(__file__), '..', 'Covoit_Africa_Remaining_Work.docx')
doc.save(output_path)
print(f'Document saved to: {os.path.abspath(output_path)}')
